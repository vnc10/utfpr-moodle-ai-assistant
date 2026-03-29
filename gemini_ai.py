import mimetypes
import os
import re
import sys
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from google import genai
    from google.genai import types
except ImportError:
    print("A biblioteca google-genai nao esta instalada.")
    print("Execute: pip install google-genai")
    sys.exit(1)

from config import (
    GEMINI_API_KEY,
    GEMINI_MODEL,
    GEMINI_MODEL_PRO,
    TEACHER_FOLDERS_DEFAULT,
    TEACHER_FOLDERS_BY_COURSE,
    TEXT_CODE_EXTENSIONS,
    MEDIA_EXTENSIONS,
    MAX_RETRIES,
    RETRY_DELAY_SECONDS,
)

client = genai.Client(api_key=GEMINI_API_KEY)


def upload_file(filepath):
    """Faz upload de um arquivo para o Gemini e aguarda processamento."""
    mime_type = mimetypes.guess_type(filepath)[0] or "application/octet-stream"
    with open(filepath, "rb") as f:
        g_file = client.files.upload(
            file=f,
            config=types.UploadFileConfig(
                display_name=os.path.basename(filepath),
                mime_type=mime_type,
            ),
        )
    while g_file.state.name == "PROCESSING":
        time.sleep(2)
        g_file = client.files.get(name=g_file.name)
    return g_file


def delete_files(file_list):
    """Remove uma lista de arquivos da nuvem do Gemini."""
    for g_file in file_list:
        try:
            client.files.delete(name=g_file.name)
        except Exception:
            pass


def _get_teacher_folders(course_name):
    """Retorna as pastas permitidas para contexto com base na disciplina."""
    name_lower = course_name.lower()
    for keyword, folders in TEACHER_FOLDERS_BY_COURSE.items():
        if keyword in name_lower:
            return folders
    return TEACHER_FOLDERS_DEFAULT


def load_teacher_context(course_name):
    """Faz upload dos PDFs do professor (slides/atividades) como contexto."""
    context_files = []
    allowed_folders = _get_teacher_folders(course_name)
    print(f"\n[AI CONTEXT] Scanning '{course_name}' for teacher PDFs...")
    print(f"   Folders filter: {allowed_folders}")

    for root, dirs, files in os.walk(course_name):
        if "Submissions_" in root:
            continue

        folder_name = os.path.basename(root)
        is_allowed = any(name in folder_name for name in allowed_folders)
        if not is_allowed and folder_name != course_name:
            continue

        for f in files:
            if not f.lower().endswith(".pdf"):
                continue
            filepath = os.path.join(root, f)
            print(f"   -> Uploading {f} to Gemini...")
            try:
                context_files.append(upload_file(filepath))
            except Exception as e:
                print(f"      [WARNING] Could not upload {f}: {e}")

    print(f"Context loaded: {len(context_files)} files uploaded.")
    return context_files


def list_slide_pdfs(course_name):
    """Lists all slide PDFs in the course directory and returns their paths."""
    slides = []
    allowed_folders = _get_teacher_folders(course_name)

    for root, dirs, files in os.walk(course_name):
        if "Submissions_" in root or "Roteiros" in root:
            continue

        folder_name = os.path.basename(root)
        is_allowed = any(name in folder_name for name in allowed_folders)
        if not is_allowed and folder_name != course_name:
            continue

        for f in sorted(files):
            if f.lower().endswith(".pdf"):
                slides.append(os.path.join(root, f))

    return slides


def upload_single_slide(filepath):
    """Uploads a single slide PDF to Gemini and returns the file reference."""
    print(f"\n[AI CONTEXT] Uploading slide: {os.path.basename(filepath)}...")
    try:
        return upload_file(filepath)
    except Exception as e:
        print(f"   [WARNING] Could not upload {os.path.basename(filepath)}: {e}")
        return None


def find_all_lesson_plans(course_name):
    """Finds all .docx lesson plans across all Roteiros folders, deduplicated by filename."""
    seen = {}
    for root, dirs, files in os.walk(course_name):
        if os.path.basename(root) != "Roteiros":
            continue
        for f in sorted(files):
            if f.lower().endswith(".docx") and f not in seen:
                seen[f] = os.path.join(root, f)
    return sorted(seen.values(), key=lambda p: os.path.basename(p))


def _read_docx_as_text(filepath):
    """Reads a .docx file and returns its content as plain text."""
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    return "\n".join(lines)


def load_previous_lesson_plans(course_name):
    """Reads all previous .docx lesson plans and returns their text content."""
    plans = find_all_lesson_plans(course_name)

    if not plans:
        return ""

    all_text = ""
    print(f"\n[AI CONTEXT] Loading {len(plans)} previous lesson plan(s)...")
    for filepath in plans:
        filename = os.path.basename(filepath)
        print(f"   -> Reading {filename}...")
        try:
            text = _read_docx_as_text(filepath)
            all_text += f"\n\n--- ROTEIRO ANTERIOR: {filename} ---\n{text}"
        except Exception as e:
            print(f"      [WARNING] Could not read {filename}: {e}")

    return all_text


def read_student_files(student_dir):
    """Le os arquivos do aluno e retorna texto e uploads para o Gemini."""
    student_text = ""
    uploaded_files = []

    for root, dirs, files in os.walk(student_dir):
        for filename in files:
            filepath = os.path.join(root, filename)
            ext = filename.lower()

            if any(ext.endswith(e) for e in TEXT_CODE_EXTENSIONS):
                try:
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        student_text += f"\nFILE: {filename}\n{f.read()}\n"
                except Exception:
                    pass

            elif any(ext.endswith(e) for e in MEDIA_EXTENSIONS):
                print(f"      [UPLOADING TO GEMINI] {filename}...")
                try:
                    uploaded_files.append(upload_file(filepath))
                except Exception as e:
                    print(f"      [WARNING] Could not upload {filename}: {e}")

    return student_text, uploaded_files


def _get_course_context(course_name):
    """Retorna a descricao do publico-alvo com base no nome da disciplina."""
    name_lower = course_name.lower()
    if "redes de computadores" in name_lower:
        return (
            "para alunos do curso Tecnico em Informatica integrado ao Ensino Medio"
        )
    if "estrutura de dados" in name_lower or "programação estruturada" in name_lower:
        return (
            "para alunos do curso superior de Engenharia Eletronica"
        )
    if "informática" in name_lower and "alimentos" in name_lower:
        return (
            "para alunos do curso superior Tecnico em Alimentos"
        )
    return "para alunos universitarios"


def generate_feedback(course_name, assignment_desc, student_text, student_files, context_files):
    """Gera feedback via Gemini com retry automatico para erros de quota."""
    public = _get_course_context(course_name)

    system_instruction = (
        f"Você é um professor universitário formado em Ciência da Computação "
        f"lecionando a disciplina de '{course_name}' {public}.\n"
        f"A NOTA DO ALUNO SERÁ SEMPRE 10, independentemente da qualidade da entrega.\n"
        f"Sua função é gerar um feedback amigável, educado e construtivo em português.\n"
        f"Utilize os PDFs de contexto fornecidos para manter a terminologia usada em sala de aula.\n\n"
        f"ATENÇÃO - REGRAS ESTRITAS DE FORMATAÇÃO:\n"
        f"1. A resposta DEVE ser escrita INTEIRAMENTE em HTML.\n"
        f"2. É ESTRITAMENTE PROIBIDO o uso de Markdown. NÃO use asteriscos (**) para negrito, use a tag <strong>. NÃO use hifens ou asteriscos (*) para listas, use <ul> e <li>.\n"
        f"3. Use a tag <br> ou <p> para quebras de linha.\n"
        f"4. NÃO envolva a resposta em blocos de código (como ```html), retorne o HTML diretamente.\n"
        f"5. OBRIGATÓRIO: No final do feedback, insira exatamente esta linha de código HTML para informar sobre o uso de IA: <br><br><hr><p><em>* Este feedback foi gerado automaticamente por IA com base no material da disciplina.</em></p>\n\n"
        f"Por fim, não cite nomes próprios. Ao se identificar no final do feedback (logo antes do aviso de IA), assine apenas como 'Professor'."
    )

    prompt = f"ENUNCIADO DA TAREFA:\n{assignment_desc}\n\nCONTEUDO DO ALUNO:\n{student_text}"
    contents = [prompt] + context_files + student_files

    for attempt in range(MAX_RETRIES):
        try:
            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=contents,
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.2,
                ),
            )

            feedback = response.text.replace("```html", "").replace("```", "").strip()
            return feedback

        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                print(
                    f"      [QUOTA EXCEEDED] Tentativa {attempt + 1}/{MAX_RETRIES}. "
                    f"Esperando {RETRY_DELAY_SECONDS}s..."
                )
                time.sleep(RETRY_DELAY_SECONDS)
            else:
                print(f"      [AI ERROR] Failed to generate feedback: {e}")
                break

    return None


def _add_formatted_run(paragraph, text):
    """Adds a run to a paragraph, handling **bold** markers."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text, output_path):
    """Converts Markdown text to a .docx file."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    for line in markdown_text.split('\n'):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph('')
            continue

        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_run(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', stripped)
            _add_formatted_run(p, text)
        elif stripped == '---':
            doc.add_paragraph('').runs
            doc.add_paragraph('_' * 50)
        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)

    doc.save(output_path)


def generate_lesson_plan(course_name, slide_file, previous_plans):
    """Generates a lesson plan via Gemini based on a slide and previous plans."""
    public = _get_course_context(course_name)

    has_previous = len(previous_plans) > 0

    system_instruction = f"""
    Você é um Professor Universitário de Ciência da Computação experiente e altamente didático, lecionando a disciplina de '{course_name}' {public}.

    Sua tarefa é analisar os slides da aula atual (em PDF) e o roteiro das aulas anteriores, para criar um roteiro de apresentação (Lesson Plan) contínuo, extremamente detalhado e pronto para ser usado como documento de texto (Google Docs/Word).

    REGRAS DE CONTINUIDADE:
    Utilize o histórico das aulas anteriores para manter a continuidade lógica. Logo no início do documento, crie um "Gancho/Recapitulação" de 1 a 2 minutos conectando os conceitos que os alunos já aprenderam com o que verão hoje.

    ESTRUTURA OBRIGATÓRIA (Use Markdown limpo):
    Gere o roteiro abordando TODOS os slides. Para cada slide, utilize exatamente esta estrutura:

    ## Slide [Número]: [Título ou Tema do Slide]
    * **Tempo Estimado:** [Quantos minutos gastar neste slide]
    * **Roteiro de Fala:** [Escreva em discurso direto, como se fosse a transcrição do que você vai falar. Não resuma: transforme os tópicos do slide em uma explicação verbal fluida e rica. Use analogias da área de tecnologia/programação para facilitar conceitos complexos.]
    * **Foco Visual/Código:** [Indique exatamente o que apontar no slide. Ex: "Chame a atenção para a linha 15 do código" ou "Explique o fluxo de dados no diagrama à direita".]
    * **Interação:** [Uma pergunta técnica rápida, teste de mesa mental ou provocação para manter a turma engajada.]

    RESTRIÇÕES:
    1. Não use tags HTML (<p>, <div>, etc). Formate tudo apenas com Markdown básico.
    2. Seja exaustivo no "Roteiro de Fala". É um guia para o professor ler e se guiar, não um mero resumo do que está escrito no slide.
    """

    previous_section = ""
    if has_previous:
        previous_section = f"\n\nROTEIROS DAS AULAS ANTERIORES:\n{previous_plans}"

    prompt = f"Analise o PDF em anexo e gere o roteiro da aula de hoje seguindo as instruções.{previous_section}"
    contents = [prompt, slide_file]

    for attempt in range(MAX_RETRIES):
        try:
            response = client.models.generate_content(
                model=GEMINI_MODEL_PRO,
                contents=contents,
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.4,
                ),
            )

            return response.text.replace("```markdown", "").replace("```", "").strip()

        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                print(
                    f"      [QUOTA EXCEEDED] Attempt {attempt + 1}/{MAX_RETRIES}. "
                    f"Waiting {RETRY_DELAY_SECONDS}s..."
                )
                time.sleep(RETRY_DELAY_SECONDS)
            else:
                print(f"      [AI ERROR] Failed to generate lesson plan: {e}")
                break

    return None
